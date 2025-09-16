import os
import tempfile
import subprocess
import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx2pdf import convert as docx2pdf_convert

def write_temp_file(content: bytes, suffix: str) -> str:
    """
    Write bytes to a uniquely-named temp file.
    Returns its full path.
    """
    tf = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
    tf.write(content)
    tf.close()
    return tf.name

def pdf_to_docx_pymupdf(input_path: str, original_name: str) -> str:
    """
    Convert PDF to DOCX using PyMuPDF (fitz) for better text extraction.
    """
    try:
        import fitz  # PyMuPDF
        
        out_dir = tempfile.gettempdir()
        base = os.path.splitext(os.path.basename(original_name))[0]
        out_path = os.path.join(out_dir, f"{base}.docx")
        
        # Open PDF
        pdf_document = fitz.open(input_path)
        doc = Document()
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Extract text blocks with positioning info
            blocks = page.get_text("dict")
            
            for block in blocks["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        line_text = ""
                        for span in line["spans"]:
                            line_text += span["text"]
                        
                        line_text = line_text.strip()
                        if line_text:
                            # Detect bullet points by indentation and common patterns
                            is_bullet = False
                            clean_text = line_text
                            
                            # Check for indentation (x coordinate)
                            x_coord = line["spans"][0]["bbox"][0] if line["spans"] else 0
                            
                            # If significantly indented, likely a bullet point
                            if x_coord > 50:  # Adjust threshold as needed
                                is_bullet = True
                                
                            # Also check for bullet-like patterns
                            bullet_patterns = [
                                r'^[\s]*[•\-\*\+○●◦‣▪▫■□▸▹▶▷][\s]*',
                                r'^[\s]*[\d]+[\.\)][\s]*',
                                r'^[\s]*[a-zA-Z][\.\)][\s]*',
                            ]
                            
                            for pattern in bullet_patterns:
                                if re.match(pattern, line_text):
                                    is_bullet = True
                                    clean_text = re.sub(pattern, '', line_text).strip()
                                    break
                            
                            # Add paragraph
                            if clean_text:
                                p = doc.add_paragraph()
                                p.add_run(clean_text)
                                
                                if is_bullet:
                                    try:
                                        p.style = 'List Bullet'
                                    except:
                                        p.paragraph_format.left_indent = Inches(0.5)
                                        p.paragraph_format.first_line_indent = Inches(-0.25)
        
        pdf_document.close()
        doc.save(out_path)
        return out_path
        
    except ImportError:
        print("PyMuPDF not available, falling back to alternative method")
        return pdf_to_docx_alternative(input_path, original_name)
    except Exception as e:
        print(f"Error with PyMuPDF: {e}")
        return pdf_to_docx_alternative(input_path, original_name)

def pdf_to_docx_alternative(input_path: str, original_name: str) -> str:
    """
    Alternative method using pdfplumber with better bullet detection.
    """
    try:
        import pdfplumber
        
        out_dir = tempfile.gettempdir()
        base = os.path.splitext(os.path.basename(original_name))[0]
        out_path = os.path.join(out_dir, f"{base}.docx")
        
        doc = Document()
        
        with pdfplumber.open(input_path) as pdf:
            for page in pdf.pages:
                # Extract text with character-level detail
                chars = page.chars
                
                # Group characters into lines
                lines = []
                current_line = []
                current_y = None
                
                for char in chars:
                    if current_y is None:
                        current_y = char['y0']
                    
                    # If y position changed significantly, it's a new line
                    if abs(char['y0'] - current_y) > 3:  # Adjust threshold as needed
                        if current_line:
                            lines.append(current_line)
                        current_line = [char]
                        current_y = char['y0']
                    else:
                        current_line.append(char)
                
                # Add the last line
                if current_line:
                    lines.append(current_line)
                
                # Process each line
                for line_chars in lines:
                    if not line_chars:
                        continue
                    
                    # Build text and check for bullets
                    line_text = ''.join(char['text'] for char in line_chars)
                    line_text = line_text.strip()
                    
                    if not line_text:
                        continue
                    
                    # Check indentation
                    left_x = min(char['x0'] for char in line_chars)
                    is_bullet = left_x > 50  # Adjust threshold
                    
                    # Check first character for bullet symbols
                    first_char = line_chars[0]['text']
                    if first_char in ['•', '▪', '▫', '■', '□', '◦', '‣', '-', '*', '+']:
                        is_bullet = True
                        # Remove the bullet character
                        line_text = line_text[1:].strip()
                    
                    # Add to document
                    if line_text:
                        p = doc.add_paragraph()
                        p.add_run(line_text)
                        
                        if is_bullet:
                            try:
                                p.style = 'List Bullet'
                            except:
                                p.paragraph_format.left_indent = Inches(0.5)
                                p.paragraph_format.first_line_indent = Inches(-0.25)
        
        doc.save(out_path)
        return out_path
        
    except ImportError:
        print("pdfplumber not available, falling back to pdf2docx")
        return pdf_to_docx_fallback(input_path, original_name)
    except Exception as e:
        print(f"Error with pdfplumber: {e}")
        return pdf_to_docx_fallback(input_path, original_name)

def pdf_to_docx_fallback(input_path: str, original_name: str) -> str:
    """
    Fallback using pdf2docx with maximum post-processing.
    """
    try:
        from pdf2docx import Converter
        
        out_dir = tempfile.gettempdir()
        base = os.path.splitext(os.path.basename(original_name))[0]
        out_path = os.path.join(out_dir, f"{base}.docx")
        
        # Convert with pdf2docx
        cv = Converter(input_path)
        cv.convert(out_path, start=0, end=None)
        cv.close()
        
        # Heavy post-processing
        fix_bullets_comprehensive(out_path)
        
        return out_path
        
    except Exception as e:
        print(f"All conversion methods failed: {e}")
        raise

def fix_bullets_comprehensive(docx_path: str):
    """
    Comprehensive bullet fixing as last resort.
    """
    doc = Document(docx_path)
    
    for paragraph in doc.paragraphs:
        text = paragraph.text
        if not text.strip():
            continue
        
        # Check every single character for problematic ones
        fixed_text = ""
        for char in text:
            char_code = ord(char)
            
            # Replace any character that's likely a problematic bullet
            if (char_code >= 0xF000 and char_code <= 0xF8FF) or \
               (char_code >= 0xE000 and char_code <= 0xF8FF) or \
               char in ['□', '■', '▪', '▫', '▬', '▭', '▮', '▯', '⬜', '⬛'] or \
               char_code in [0x2022, 0x25CF, 0x2219, 0x2043, 0x204C]:
                # If it's the first character, replace with bullet
                if len(fixed_text) == 0:
                    fixed_text = "• "
                else:
                    fixed_text += "•"
            else:
                fixed_text += char
        
        # If text was modified, update the paragraph
        if fixed_text != text:
            paragraph.clear()
            paragraph.add_run(fixed_text.strip())
            
            # Apply list formatting if it starts with bullet
            if fixed_text.strip().startswith('•'):
                try:
                    paragraph.style = 'List Bullet'
                except:
                    paragraph.paragraph_format.left_indent = Inches(0.5)
                    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    
    doc.save(docx_path)

def pdf_to_docx(input_path: str, original_name: str) -> str:
    """
    Main conversion function that tries multiple methods.
    """
    # Try PyMuPDF first (most reliable)
    try:
        return pdf_to_docx_pymupdf(input_path, original_name)
    except:
        pass
    
    # Try pdfplumber second
    try:
        return pdf_to_docx_alternative(input_path, original_name)
    except:
        pass
    
    # Fall back to pdf2docx with heavy post-processing
    return pdf_to_docx_fallback(input_path, original_name)

def docx_to_pdf(input_path: str, original_name: str) -> str:
    """
    Convert DOCX to PDF.
    """
    out_dir = tempfile.gettempdir()
    base = os.path.splitext(os.path.basename(original_name))[0]
    out_path = os.path.join(out_dir, f"{base}.pdf")

    # Try docx2pdf first
    try:
        docx2pdf_convert(input_path, out_path)
        if os.path.exists(out_path):
            return out_path
    except Exception:
        pass

    # If docx2pdf fails, raise an error
    raise Exception("Could not convert DOCX to PDF. Please ensure docx2pdf is working.")